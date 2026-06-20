from pathlib import Path
import textwrap

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT = Path(__file__).resolve().parent
SCREENSHOT_DIR = ROOT / "report_screenshots"
ASSET_DIR = ROOT / "_report_assets"
OUT_DOCX = ROOT / "国产数据库开发与应用说明报告_待替换截图.docx"


FIGURES = [
    {
        "key": "fig01",
        "file": "fig01_kingbase_tables.png",
        "caption": "图 1 KingbaseES 中创建 4 张业务数据表",
        "title": "KingbaseES 中创建 4 张业务数据表",
        "requirements": [
            "显示数据库 course_assets 或当前连接库",
            "能看到 kb_department、kb_asset_category、kb_supplier、kb_asset 四张表",
            "优先展示数据库工具的对象树、表列表或查询 information_schema 的结果",
        ],
        "width_cm": 13.2,
    },
    {
        "key": "fig02",
        "file": "fig02_project_startup.png",
        "caption": "图 2 Spring Boot 项目启动并连接金仓数据库",
        "title": "Spring Boot 项目启动并连接金仓数据库",
        "requirements": [
            "终端执行 mvn spring-boot:run 或 java -jar",
            "能看到 Tomcat started on port 8080 或 Started KingbaseAssetApplication",
            "如出现 schema.sql/data.sql 初始化日志，可一并截入",
        ],
        "width_cm": 13.2,
    },
    {
        "key": "fig03",
        "file": "fig03_asset_list.png",
        "caption": "图 3 教学设备资产列表页面",
        "title": "教学设备资产列表页面",
        "requirements": [
            "浏览器打开 http://localhost:8080/assets",
            "页面能看到资产编号、资产名称、分类、使用部门、供应商、状态等列",
            "至少显示两条演示数据或真实录入数据",
        ],
        "width_cm": 13.4,
    },
    {
        "key": "fig04",
        "file": "fig04_asset_search.png",
        "caption": "图 4 按条件查询资产记录",
        "title": "按条件查询资产记录",
        "requirements": [
            "在资产列表页输入关键字或选择分类、状态",
            "点击查询后列表只显示符合条件的数据",
            "截图中要能看到查询条件和筛选后的结果",
        ],
        "width_cm": 13.4,
    },
    {
        "key": "fig05",
        "file": "fig05_asset_create.png",
        "caption": "图 5 新增教学设备资产",
        "title": "新增教学设备资产",
        "requirements": [
            "点击新增资产进入表单页面",
            "表单中填写资产编号、名称、分类、部门、供应商、状态等字段",
            "建议截提交前的完整表单，或提交成功后带成功提示的列表页",
        ],
        "width_cm": 13.4,
    },
    {
        "key": "fig06",
        "file": "fig06_asset_update.png",
        "caption": "图 6 修改资产信息",
        "title": "修改资产信息",
        "requirements": [
            "在列表中点击某条资产的修改按钮",
            "编辑页面显示原资产信息，并对状态、保管人或备注进行修改",
            "截图中应能看出这是修改页面，不是新增页面",
        ],
        "width_cm": 13.4,
    },
    {
        "key": "fig07",
        "file": "fig07_asset_delete.png",
        "caption": "图 7 删除资产后列表更新",
        "title": "删除资产后列表更新",
        "requirements": [
            "点击删除按钮并确认",
            "列表返回后显示删除成功提示或目标记录已消失",
            "建议截图删除后的列表页面，不要只截浏览器确认弹窗",
        ],
        "width_cm": 13.4,
    },
    {
        "key": "fig08",
        "file": "fig08_database_result.png",
        "caption": "图 8 数据库中资产数据与 Web 操作保持一致",
        "title": "数据库中资产数据与 Web 操作保持一致",
        "requirements": [
            "在 KingbaseES 查询工具中执行 SELECT 查询 kb_asset",
            "能看到 Web 端新增或修改后的资产记录",
            "建议截 SQL 语句和查询结果区域",
        ],
        "width_cm": 13.2,
    },
]


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def font_path(name: str) -> str | None:
    path = Path("C:/Windows/Fonts") / name
    return str(path) if path.exists() else None


def load_font(size: int, bold: bool = False):
    for name in (["msyhbd.ttc", "simhei.ttf", "msyh.ttc", "simsun.ttc"] if bold else ["msyh.ttc", "simsun.ttc", "simhei.ttf"]):
        path = font_path(name)
        if path:
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_text(text: str, width: int) -> list[str]:
    return textwrap.wrap(text, width=width, break_long_words=False, replace_whitespace=False)


def create_placeholder(fig: dict) -> Path:
    ensure_dir(ASSET_DIR)
    out = ASSET_DIR / f"{fig['key']}_placeholder.png"
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    border = (144, 152, 160)
    title_color = (24, 76, 72)
    fill = (243, 247, 247)
    draw.rounded_rectangle((52, 52, 1548, 848), radius=8, fill=fill, outline=border, width=3)
    draw.rectangle((52, 52, 1548, 166), fill=(226, 237, 235))
    draw.line((52, 166, 1548, 166), fill=border, width=2)

    title_font = load_font(42, True)
    subtitle_font = load_font(27, True)
    body_font = load_font(28)
    small_font = load_font(23)

    draw.text((92, 84), "待替换为队员实操截图", font=title_font, fill=title_color)
    draw.text((92, 206), fig["caption"], font=subtitle_font, fill=(18, 24, 32))
    draw.text((92, 256), f"建议文件名：report_screenshots/{fig['file']}", font=small_font, fill=(92, 99, 105))
    draw.text((92, 330), "截图必须包含：", font=subtitle_font, fill=(18, 24, 32))
    y = 386
    for i, req in enumerate(fig["requirements"], start=1):
        for line in wrap_text(f"{i}. {req}", 42):
            draw.text((120, y), line, font=body_font, fill=(30, 36, 42))
            y += 43
        y += 8
    note = "说明：该图仅用于占位，最终提交前请用真实截图替换。"
    draw.text((92, 760), note, font=small_font, fill=(126, 74, 0))
    img.save(out)
    return out


def screenshot_path(fig: dict) -> Path:
    if SCREENSHOT_DIR.exists():
        primary = SCREENSHOT_DIR / fig["file"]
        if primary.exists():
            return primary
        stem = Path(fig["file"]).stem
        for ext in (".png", ".jpg", ".jpeg"):
            candidate = SCREENSHOT_DIR / f"{stem}{ext}"
            if candidate.exists():
                return candidate
    return create_placeholder(fig)


def normalize_image(path: Path, key: str) -> Path:
    ensure_dir(ASSET_DIR)
    out = ASSET_DIR / f"{key}_normalized.png"
    with Image.open(path) as img:
        img = ImageOps.exif_transpose(img).convert("RGB")
        img.thumbnail((1800, 1400), Image.LANCZOS)
        canvas = Image.new("RGB", img.size, "white")
        canvas.paste(img, (0, 0))
        canvas.save(out, optimize=True)
    return out


def set_run_font(run, size=None, bold=False):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_style_font(style, size, bold=False):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def set_keep_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def add_paragraph(doc, text="", style="Body", indent=True, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Pt(24) if indent else None
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 12)
    return p


def add_heading(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    set_keep_next(p)
    run = p.add_run(text)
    set_run_font(run, 15 if level == 1 else 13, True)
    return p


def add_meta(doc, label, value):
    p = doc.add_paragraph(style="Meta")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label)
    set_run_font(r1, 11, True)
    r2 = p.add_run(value)
    set_run_font(r2, 11)


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        set_run_font(run, 10.5)


def add_caption(doc, caption):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(caption)
    set_run_font(run, 10.5)


def add_figure(doc, fig):
    img = normalize_image(screenshot_path(fig), fig["key"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(img), width=Cm(fig["width_cm"]))
    add_caption(doc, fig["caption"])


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    styles = doc.styles
    set_style_font(styles["Normal"], 12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(6)
    body = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(body, 12)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(6)
    meta = styles.add_style("Meta", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(meta, 11)
    code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, 10.5)
    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Caption", 10.5)]:
        set_style_font(styles[name], size, name.startswith("Heading"))
    return doc


def build_doc() -> Path:
    doc = setup_doc()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("国产数据库（金仓数据库）开发与应用报告")
    set_run_font(run, 18, True)

    add_meta(doc, "课程主题：", "金仓数据库开发与应用")
    add_meta(doc, "项目名称：", "教学设备资产管理系统")
    add_meta(doc, "数据库：", "KingbaseES（金仓数据库）")
    add_meta(doc, "实现方式：", "Spring Boot + JdbcTemplate + Thymeleaf + KingbaseES JDBC")
    add_paragraph(
        doc,
        "本报告围绕“国产数据库（金仓数据库）开发项目”作业要求展开，说明如何基于 KingbaseES 完成一个可运行的 Web 业务系统。项目选择教学设备资产管理作为业务场景，围绕设备资产的登记、查询、维护和删除展开，实现了数据库建模、后端数据访问、Web 页面展示和增删改查操作的完整流程。",
    )

    add_heading(doc, "一、作业要求与完成情况")
    add_paragraph(
        doc,
        "本次作业要求基于国产数据库（金仓数据库）开发项目，在金仓数据库上建立 3 个以上数据表，并在 Web 端或移动终端至少完成一个与所选项目相关的业务功能，业务功能需要包含增删改查。本项目围绕教学设备资产管理系统进行设计，数据库采用 KingbaseES，建立了部门、设备分类、供应商和资产 4 张数据表，Web 端围绕资产台账完成新增、查询、修改和删除操作。",
    )
    add_paragraph(
        doc,
        "从完成情况看，系统不仅满足“3 个以上数据表”的基本要求，还通过外键把资产表与部门表、分类表、供应商表关联起来，使业务数据具有较完整的组织关系。Web 页面能够展示资产列表、按条件筛选资产、新增设备、编辑设备信息并删除记录，能够体现国产数据库在实际业务系统中的开发与应用过程。",
    )
    add_paragraph(
        doc,
        "关于安装方式，经本机验证，当前常用 Scoop bucket 中没有 KingbaseES 安装包，执行 scoop search kingbase 返回 No matches found。因此，金仓数据库本体不能直接通过 Scoop 安装，实际部署时应使用金仓官方安装包、官方镜像或老师提供的实验环境。Java 项目中的 KingbaseES JDBC 驱动则可以通过 Maven 依赖 cn.com.kingbase:kingbase8 引入。",
    )

    add_heading(doc, "二、系统功能流程说明")
    add_paragraph(
        doc,
        "系统的业务流程从数据库建模开始。首先在 KingbaseES 中创建 course_assets 数据库，并执行 schema.sql 建立 kb_department、kb_asset_category、kb_supplier 和 kb_asset 4 张表；随后执行 data.sql 写入部门、分类、供应商和资产演示数据；最后启动 Spring Boot 项目，通过 KingbaseES JDBC 驱动连接数据库并提供 Web 端访问入口。",
    )
    add_paragraph(
        doc,
        "用户进入资产列表页面后，可以看到设备编号、设备名称、设备分类、使用部门、供应商、资产原值、状态和保管人等字段。页面上方提供关键字、分类和状态筛选条件，用户可以根据资产编号、名称、部门或保管人快速定位数据。列表中的“新增资产”“修改”“删除”等入口组成了完整的资产台账业务流程。",
    )
    add_paragraph(
        doc,
        "新增资产时，用户在表单中填写资产编号、资产名称、设备分类、使用部门、供应商、购置日期、原值、状态、保管人和备注。系统提交后将数据写入 kb_asset 表，并通过外键保存与分类、部门和供应商的关联。修改资产时，系统先读取原有记录并回显到表单中，用户调整状态、保管人或备注后保存。删除资产时，系统根据资产 id 删除对应记录，列表随即更新。",
    )

    add_heading(doc, "三、系统设计与实现要点")
    add_heading(doc, "3.1 数据库表结构设计", level=2)
    add_paragraph(
        doc,
        "数据库设计采用“基础字典表 + 主业务表”的方式。kb_department 保存设备使用部门，kb_asset_category 保存设备类别，kb_supplier 保存供应商信息，kb_asset 保存具体资产记录。资产表通过 category_id、department_id 和 supplier_id 分别关联分类、部门和供应商，既减少了重复数据，也便于后续按照分类、部门或供应商进行统计。",
    )
    add_code(doc, [
        "kb_department(id, name, manager, phone, created_at)",
        "kb_asset_category(id, name, description, created_at)",
        "kb_supplier(id, name, contact_person, phone, address, created_at)",
        "kb_asset(id, asset_code, asset_name, category_id, department_id, supplier_id, purchase_date, original_value, status, keeper, remark)",
    ])
    add_paragraph(
        doc,
        "kb_asset 表中的 asset_code 设置唯一约束，防止同一设备编号重复录入；status 字段使用检查约束限制在 IN_STOCK、IN_USE、REPAIR 和 SCRAPPED 四种状态内，避免页面误传非法状态。通过这些约束，数据库不仅负责存储数据，也参与保证业务数据的准确性。",
    )

    add_heading(doc, "3.2 后端连接与数据访问", level=2)
    add_paragraph(
        doc,
        "后端使用 Spring Boot 构建，数据访问层采用 JdbcTemplate，数据库驱动配置为 com.kingbase8.Driver，连接地址采用 jdbc:kingbase8://localhost:54321/course_assets。相比只写静态页面，本项目真正通过后端读取和写入 KingbaseES 数据表，能够体现国产数据库在 Web 项目中的实际使用方式。",
    )
    add_code(doc, [
        "spring.datasource.driver-class-name=com.kingbase8.Driver",
        "spring.datasource.url=jdbc:kingbase8://localhost:54321/course_assets",
        "spring.datasource.username=system",
        "spring.datasource.password=manager",
    ])
    add_paragraph(
        doc,
        "AssetRepository 中封装了资产查询、新增、修改和删除 SQL。查询方法通过 JOIN 同时读取资产表、分类表、部门表和供应商表，使页面能够直接展示中文业务名称；新增和修改方法通过表单对象接收页面字段；删除方法根据 id 删除资产记录。Controller 层负责接收 Web 请求、进行简单校验并返回 Thymeleaf 页面。",
    )

    add_heading(doc, "3.3 Web 页面与交互设计", level=2)
    add_paragraph(
        doc,
        "Web 端使用 Thymeleaf 模板渲染页面，主要包括资产列表页面和资产表单页面。列表页面承担查询和数据展示，表单页面同时服务新增和修改。页面布局采用左侧导航和右侧业务区域的结构，字段排列清楚，按钮位置符合常规业务系统使用习惯，便于用户连续完成查询、录入、维护和删除操作。",
    )
    add_paragraph(
        doc,
        "为了让作业功能更完整，页面提供了关键字、分类和状态三个查询条件。新增或修改失败时，系统会返回表单并提示必填项或资产编号重复等错误；删除操作增加确认提示，避免用户误删资产记录。这些细节使系统更接近真实业务应用，而不是只完成单条 SQL 演示。",
    )

    add_heading(doc, "四、执行结果截图与功能说明")
    add_paragraph(
        doc,
        "本节截图按照数据库建表、项目启动、列表查询、新增、修改、删除和数据库结果核对的顺序排列。最终提交前需要将占位图替换为队员实际运行 KingbaseES 和 Web 系统后的截图，截图文件名与 picture_capture.md 中的要求保持一致。",
    )
    descriptions = [
        "在 KingbaseES 中执行建表脚本后，可以看到系统创建了部门、设备分类、供应商和资产 4 张数据表。资产表与其他三张基础表建立外键关系，能够支撑教学设备资产登记和维护业务。",
        "启动 Spring Boot 项目后，控制台显示应用已经运行在 8080 端口，并通过 KingbaseES JDBC 驱动连接数据库。项目启动时会执行初始化脚本，确保表结构和演示数据可用于页面测试。",
        "浏览器访问资产列表页面后，系统展示资产台账数据。页面中可以看到资产编号、资产名称、分类、使用部门、供应商、原值、状态和保管人等字段，说明 Web 端已经从 KingbaseES 中读取业务数据。",
        "在列表页面输入关键字或选择分类、状态后点击查询，系统根据条件重新检索资产记录。筛选后的列表能够帮助用户快速定位某类设备或某个状态下的资产。",
        "新增资产页面用于录入新的教学设备。用户填写资产编号、名称、分类、部门、供应商等信息后提交，系统会将数据写入 kb_asset 表，并返回资产列表页面。",
        "修改资产页面会回显原有资产信息。用户可以调整资产状态、保管人或备注等字段，保存后列表中的对应记录同步更新，体现了业务数据维护过程。",
        "删除资产时，用户在列表中点击删除并确认，系统删除对应记录后返回列表页面。删除后的列表不再显示该资产，说明 Web 端删除操作已经作用到数据库。",
        "在 KingbaseES 查询工具中执行 SELECT 查询，可以看到 Web 端新增或修改后的资产记录。数据库结果与页面操作保持一致，说明前端、后端和金仓数据库之间的数据链路完整。",
    ]
    for fig, desc in zip(FIGURES, descriptions):
        add_paragraph(doc, desc)
        add_figure(doc, fig)

    add_heading(doc, "五、测试与运行说明")
    add_paragraph(
        doc,
        "项目运行前需要先安装并启动 KingbaseES，创建 course_assets 数据库。如果数据库端口、用户名或密码与默认配置不同，可以修改 application.yml，也可以通过 KINGBASE_URL、KINGBASE_USERNAME 和 KINGBASE_PASSWORD 环境变量覆盖。准备完成后，在项目目录执行 mvn spring-boot:run，即可访问 http://localhost:8080/assets。",
    )
    add_paragraph(
        doc,
        "本项目已完成编译验证，使用命令 mvn -s maven-central-settings.xml -q -DskipTests package 可以成功生成构建结果。由于当前机器未直接安装 KingbaseES 服务，最终运行截图需要由已安装金仓数据库环境的队员完成。截图时应确保数据库表、Web 页面和操作结果来自同一套运行环境。",
    )

    add_heading(doc, "六、总结")
    add_paragraph(
        doc,
        "通过本次金仓数据库开发与应用实践，可以看到国产数据库能够按照常规企业 Web 项目的方式接入业务系统。KingbaseES 负责保存结构化业务数据，Spring Boot 负责接收请求并组织业务流程，JdbcTemplate 负责执行 SQL，Thymeleaf 页面负责展示和提交数据。整个系统覆盖了从数据库建模到 Web 增删改查的完整链路。",
    )
    add_paragraph(
        doc,
        "本项目围绕教学设备资产管理这一具体场景建立 4 张业务数据表，并完成资产台账的查询、新增、修改和删除功能，符合任务中“3 个以上数据表”和“Web 端业务功能含增删改查”的要求。后续如果继续扩展，可以增加登录权限、资产盘点、维修记录、报废审批和统计报表等功能，使系统从课程作业进一步发展为更完整的资产管理应用。",
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_doc())
