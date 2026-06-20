from pathlib import Path
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT = Path(__file__).resolve().parent
SCREENSHOT_DIR = ROOT / "my_screenshots"
ASSET_DIR = ROOT / "_optimized_assets"
OUT_DOCX = ROOT / "大模型技术在系统中的应用说明报告_优化版_待替换截图.docx"


FIGURES = [
    {
        "key": "fig01",
        "file": "fig01_dify_deploy.png",
        "caption": "图 1 本地部署 Dify 并初始化管理员账号",
        "title": "本地部署 Dify 并初始化管理员账号",
        "requirements": [
            "Docker Desktop 中 Dify 相关容器处于 Running 状态",
            "浏览器访问 localhost 后显示 Dify 初始化或登录页面",
            "能看到本地服务已经启动并可进入平台",
        ],
        "width_cm": 13.0,
    },
    {
        "key": "fig02",
        "file": "fig02_model_provider.png",
        "caption": "图 2 在 Dify 中关联本地 DeepSeek 模型",
        "title": "在 Dify 中关联本地 DeepSeek 模型",
        "requirements": [
            "进入设置中的模型供应商页面",
            "显示 Ollama 或本地模型供应商配置",
            "能看到 deepseek-r1:1.5b 或实际使用的模型名称",
        ],
        "width_cm": 13.0,
    },
    {
        "key": "fig03",
        "file": "fig03_chat_app_config.png",
        "caption": "图 3 创建聊天助手并配置模型参数",
        "title": "创建聊天助手并配置模型参数",
        "requirements": [
            "工作室中已创建聊天助手应用",
            "页面显示应用名称、模型选择和参数设置",
            "能体现用户问题将由本地模型处理",
        ],
        "width_cm": 13.1,
    },
    {
        "key": "fig04",
        "file": "fig04_knowledge_indexing.png",
        "caption": "图 4 上传资料并完成知识库向量化配置",
        "title": "上传资料并完成知识库向量化配置",
        "requirements": [
            "知识库中上传了自己准备的课程或业务资料",
            "页面显示分段、Embedding、索引或处理完成状态",
            "不要使用 PDF 附件中的页面截图",
        ],
        "width_cm": 13.1,
    },
    {
        "key": "fig05",
        "file": "fig05_rag_chat_test.png",
        "caption": "图 5 将知识库添加到聊天应用并进行问答测试",
        "title": "将知识库添加到聊天应用并进行问答测试",
        "requirements": [
            "聊天应用已经关联知识库",
            "用户输入与资料相关的问题",
            "回答区域显示模型结合知识库生成的回复",
        ],
        "width_cm": 13.1,
    },
    {
        "key": "fig06",
        "file": "fig06_sql_workflow_nodes.png",
        "caption": "图 6 SQL 查询工作流的节点编排",
        "title": "SQL 查询工作流的节点编排",
        "requirements": [
            "工作流画布包含开始、LLM、参数提取器、SQL 查询等节点",
            "节点之间有清晰连接线",
            "能看出自然语言输入被转成数据库查询流程",
        ],
        "width_cm": 13.2,
    },
    {
        "key": "fig07",
        "file": "fig07_parameter_extractor.png",
        "caption": "图 7 使用参数提取器提取 SQL 语句",
        "title": "使用参数提取器提取 SQL 语句",
        "requirements": [
            "右侧配置面板显示参数提取器节点",
            "能看到 sql 或 sql_list 参数名称",
            "提取指令要求只提取可执行 SQL 语句",
        ],
        "width_cm": 13.1,
    },
    {
        "key": "fig08",
        "file": "fig08_mcp_agent_tools.png",
        "caption": "图 8 MCP 用于连接外部工具与数据源",
        "title": "MCP 用于连接外部工具与数据源",
        "requirements": [
            "优先展示 Dify 插件、工具或 Marketplace 中的 MCP 相关页面",
            "如果没有 MCP 页面，可展示 Agent 工具配置区域",
            "能体现大模型后续可以调用外部工具或数据源",
        ],
        "width_cm": 12.8,
    },
]


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def font_path(name: str) -> str | None:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts") / name.upper(),
        Path("C:/Windows/Fonts") / name.lower(),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def load_font(size: int, bold: bool = False):
    names = ["msyhbd.ttc", "simhei.ttf", "msyh.ttc", "simsun.ttc"] if bold else ["msyh.ttc", "simsun.ttc", "simhei.ttf"]
    for name in names:
        path = font_path(name)
        if path:
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_text(text: str, width: int) -> list[str]:
    wrapped = []
    for part in text.splitlines():
        if not part:
            wrapped.append("")
            continue
        wrapped.extend(textwrap.wrap(part, width=width, break_long_words=False, replace_whitespace=False))
    return wrapped


def create_placeholder(fig: dict) -> Path:
    ensure_dir(ASSET_DIR)
    out = ASSET_DIR / f"{fig['key']}_placeholder.png"
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    border = (145, 145, 145)
    accent = (42, 91, 150)
    muted = (86, 86, 86)
    light = (244, 247, 251)

    draw.rounded_rectangle((50, 50, 1550, 850), radius=10, outline=border, width=3, fill=light)
    draw.rectangle((50, 50, 1550, 168), fill=(230, 237, 246), outline=border, width=0)
    draw.line((50, 168, 1550, 168), fill=border, width=2)

    title_font = load_font(42, bold=True)
    body_font = load_font(29)
    small_font = load_font(24)
    hint_font = load_font(26, bold=True)

    draw.text((90, 84), "待替换为本人实操截图", font=title_font, fill=accent)
    draw.text((90, 205), fig["caption"], font=hint_font, fill=(0, 0, 0))
    draw.text((90, 258), f"建议文件名：my_screenshots/{fig['file']}", font=small_font, fill=muted)

    y = 335
    draw.text((90, y), "截图必须包含：", font=hint_font, fill=(0, 0, 0))
    y += 52
    for idx, req in enumerate(fig["requirements"], 1):
        lines = wrap_text(f"{idx}. {req}", 42)
        for line in lines:
            draw.text((120, y), line, font=body_font, fill=(25, 25, 25))
            y += 43
        y += 8

    note = "说明：该占位图不是最终提交截图。请按 picture_capture.md 完成真实截图后重新生成或替换本图。"
    for line in wrap_text(note, 50):
        draw.text((90, 760), line, font=small_font, fill=(120, 56, 0))
        break

    img.save(out)
    return out


def screenshot_path(fig: dict) -> Path:
    primary = SCREENSHOT_DIR / fig["file"]
    if primary.exists():
        return primary
    lower_matches = list(SCREENSHOT_DIR.glob(fig["file"].replace(".png", ".*"))) if SCREENSHOT_DIR.exists() else []
    for match in lower_matches:
        if match.suffix.lower() in {".png", ".jpg", ".jpeg"}:
            return match
    return create_placeholder(fig)


def normalize_image(src: Path, key: str) -> Path:
    ensure_dir(ASSET_DIR)
    out = ASSET_DIR / f"{key}_normalized.png"
    with Image.open(src) as im:
        im = ImageOps.exif_transpose(im).convert("RGB")
        # Keep screenshots readable while avoiding extremely large DOCX payloads.
        max_w, max_h = 1800, 1400
        im.thumbnail((max_w, max_h), Image.LANCZOS)
        canvas = Image.new("RGB", im.size, "white")
        canvas.paste(im, (0, 0))
        canvas.save(out, optimize=True)
    return out


def set_run_font(run, size=None, bold=False, color=(0, 0, 0)):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def set_style_font(style, size, bold=False):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def set_keep_next(paragraph, enabled=True):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:keepNext"))
    if enabled and existing is None:
        p_pr.append(OxmlElement("w:keepNext"))
    if not enabled and existing is not None:
        p_pr.remove(existing)


def add_paragraph(doc, text="", style="Body", indent=True, after=6, size=12):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Pt(24) if indent else None
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_heading(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    set_keep_next(p, True)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 13, bold=True)
    return p


def add_metadata_line(doc, label, value):
    p = doc.add_paragraph(style="Meta")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(label)
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(value)
    set_run_font(r2, size=11)


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(line)
        set_run_font(r, size=10.5)


def add_caption(doc, caption):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(caption)
    set_run_font(r, size=10.5)
    return p


def add_figure(doc, fig: dict):
    image = normalize_image(screenshot_path(fig), fig["key"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image), width=Cm(fig["width_cm"]))
    add_caption(doc, fig["caption"])


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    set_style_font(styles["Normal"], 12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(6)

    body = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(body, 12)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(6)

    meta = styles.add_style("Meta", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(meta, 11)
    meta.paragraph_format.space_after = Pt(2)

    code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, 10.5)
    code.paragraph_format.line_spacing = 1.15

    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Caption", 10.5)]:
        set_style_font(styles[name], size, bold=name.startswith("Heading"))
    return doc


def build_doc() -> Path:
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 1.15
    r = title.add_run("大模型技术在系统中的应用说明报告")
    set_run_font(r, size=18, bold=True)

    add_metadata_line(doc, "课程内容：", "大模型技术培训")
    add_metadata_line(doc, "培训时间：", "2026 年 6 月 17 日 18:30")
    add_metadata_line(doc, "培训地点：", "信息 A109")
    add_metadata_line(doc, "主讲教师：", "宋航老师")

    add_paragraph(
        doc,
        "本报告围绕“大模型技术”培训任务展开，结合 Dify 平台、本地 Ollama 模型服务与知识库/工作流能力，说明大模型技术如何被接入到一个可运行、可交互、可扩展的系统中。报告重点不是停留在概念介绍，而是按照系统落地过程说明模型部署、模型调用、知识库增强问答、SQL 查询工作流和 MCP/Agent 扩展的实现思路。",
    )

    add_heading(doc, "一、作业要求与完成情况")
    add_paragraph(
        doc,
        "本次作业要求提交一个文档，说明如何在系统中使用大模型技术。围绕这一要求，本系统以 Dify 作为大模型应用开发与编排平台，以 Ollama 作为本地模型运行环境，接入 DeepSeek 系列本地模型完成基础对话、知识库问答、自然语言转 SQL 和外部工具扩展设计。Dify 负责应用创建、知识库管理、提示词配置和工作流调度，Ollama 负责提供本地模型推理服务，两者配合后形成完整的大模型应用运行链路。",
    )
    add_paragraph(
        doc,
        "从功能完成情况看，系统中大模型技术主要体现在四个层面。第一，作为自然语言交互入口，用户可以通过聊天助手直接向系统提问；第二，作为知识库问答引擎，模型结合检索到的资料片段生成更贴近课程或业务资料的回答；第三，作为流程编排中的智能节点，模型根据用户输入生成结构化 SQL 查询方案；第四，作为后续 Agent 能力的核心推理模块，模型可通过 MCP 或工具调用机制连接外部数据源和服务。",
    )
    add_paragraph(
        doc,
        "因此，本系统并不是把大模型当作一个孤立的聊天窗口，而是把大模型放入实际系统流程中，让它承担“理解问题、生成内容、组织结构化结果、辅助调用工具”的作用。知识库、数据库查询组件和外部工具则负责提供数据与执行能力，这种分工更符合真实大模型应用的建设方式。",
    )

    add_heading(doc, "二、系统功能流程说明")
    add_paragraph(
        doc,
        "系统整体流程可以概括为“部署平台—接入模型—创建应用—配置知识库—编排工作流—扩展工具能力”。首先使用 Docker 在本地部署 Dify，并通过浏览器完成管理员账号初始化；随后安装 Ollama，下载并运行 DeepSeek 本地模型；接着在 Dify 的模型供应商页面中配置 Ollama，使 Dify 应用能够调用本机模型服务；最后根据实际场景分别创建聊天助手、知识库应用和 SQL 查询工作流。",
    )
    add_paragraph(
        doc,
        "在基础对话流程中，用户在聊天助手中输入问题，Dify 将用户输入、提示词和参数配置发送给本地模型，模型完成推理后返回自然语言答案。这个流程适合课程答疑、内容解释、代码思路分析和一般性学习辅助。",
    )
    add_paragraph(
        doc,
        "在知识库问答流程中，系统先把上传的课程资料或业务文档解析成文本片段，再通过 Embedding 模型完成向量化和索引。用户提问时，Dify 先从知识库中检索与问题最相关的片段，再把检索内容连同问题交给大模型生成答案。这样可以减少模型脱离资料随意发挥的情况，也能让回答更贴近本系统掌握的本地知识。",
    )
    add_paragraph(
        doc,
        "在 SQL 查询流程中，用户不需要直接编写 SQL，而是用自然语言描述查询需求。工作流中的 LLM 节点根据数据库类型、表结构、字段含义和输出格式要求生成查询语句，参数提取器再把 SQL 从模型输出中提取成可被后续组件读取的结构化结果，最后由数据库查询节点执行。该流程体现了大模型在系统中的核心价值：用自然语言理解能力降低用户操作门槛，用工作流组件保证执行过程可控。",
    )
    add_paragraph(
        doc,
        "在扩展流程中，MCP 可以作为模型连接外部工具和数据源的标准协议。通过 MCP 或 Agent 工具调用机制，系统后续可以接入天气 API、企业数据库、本地文件系统或开发工具，使大模型从“回答问题”进一步发展为“调用工具完成任务”。",
    )

    add_heading(doc, "三、系统设计与实现要点")
    add_heading(doc, "3.1 平台部署与模型接入", level=2)
    add_paragraph(
        doc,
        "Dify 采用 Docker 部署，平台内部包含 web、api、worker、db、redis、nginx 等服务组件。Ollama 运行在本机环境中，用于启动和管理本地大模型。由于 Dify 容器需要访问宿主机上的 Ollama 服务，因此配置时需要保证容器能够访问 Ollama 的 API 地址。完成配置后，Dify 中的聊天助手、知识库和工作流都可以复用同一个模型供应商。",
    )
    add_code(
        doc,
        [
            "# 启用自定义模型能力",
            "CUSTOM_MODEL_ENABLED=true",
            "# 让 Dify 容器访问宿主机 Ollama 服务",
            "OLLAMA_API_BASE_URL=host.docker.internal:11434",
        ],
    )
    add_paragraph(
        doc,
        "模型接入完成后，需要在 Dify 的“设置—模型供应商”中添加 Ollama，并在系统模型设置中选择 DeepSeek 模型。该步骤相当于把底层推理能力注册到应用平台中，后续所有应用只需要在 Dify 中选择模型即可调用。",
    )

    add_heading(doc, "3.2 聊天助手的提示词与参数", level=2)
    add_paragraph(
        doc,
        "聊天助手是系统中最直接的大模型入口。创建应用时需要设置应用名称、描述、模型和推理参数。温度、Top P、上下文长度和最大输出长度会影响回答效果：课程问答和 SQL 生成更强调准确性与稳定性，参数应偏保守；开放式内容生成可以适当提高创造性。通过合理设置参数，系统能够在“回答稳定”和“表达自然”之间取得平衡。",
    )
    add_paragraph(
        doc,
        "提示词用于约束模型的角色、任务边界和输出格式。在本系统中，提示词不只是告诉模型“回答问题”，还会说明回答应基于知识库资料、SQL 需要符合 MySQL 5.7 语法、输出应保持 JSON 或数组格式等要求。这样可以让大模型的输出更稳定，也方便后续节点继续处理。",
    )

    add_heading(doc, "3.3 知识库增强问答", level=2)
    add_paragraph(
        doc,
        "知识库部分采用 RAG 思路。Dify 将上传的文档解析、分段并向量化，当用户提问时先进行相似度检索，再把检索结果作为上下文交给模型。模型负责理解问题和组织答案，知识库负责提供可靠资料来源。相比只调用通用模型，RAG 能让系统回答课程讲义、项目文档和业务说明中的具体问题。",
    )
    add_paragraph(
        doc,
        "为了提高问答质量，知识库资料应尽量使用结构清晰的文档，分段长度不宜过长，文件名和知识库名称也应能反映资料主题。测试时可以优先提问资料中明确出现过的问题，通过回答内容、引用片段和检索结果判断知识库是否真正发挥作用。",
    )

    add_heading(doc, "3.4 SQL 工作流与结构化输出", level=2)
    add_paragraph(
        doc,
        "SQL 查询工作流把用户输入、大模型理解、参数提取和数据库执行拆成多个节点。LLM 节点根据提示词生成 SQL，参数提取器把模型输出中的 SQL 字段提取出来，数据库查询节点再读取该字段并执行。这样的设计避免让模型直接操作数据库，而是把模型输出放入受控流程中处理。",
    )
    add_paragraph(
        doc,
        "为了提高流程安全性，实际使用时应限制模型只能生成 SELECT 查询，禁止 DROP、DELETE、UPDATE 等危险语句，并对表名、字段名和查询条数做白名单或上限控制。同时应记录用户问题、模型输出和数据库执行结果，便于后续排查错误和优化提示词。",
    )

    add_heading(doc, "3.5 MCP 与 Agent 扩展", level=2)
    add_paragraph(
        doc,
        "MCP（Model Context Protocol，模型上下文协议）用于解决大模型连接外部数据源和工具的问题。传统方式通常需要为不同模型、不同工具分别开发接口，而 MCP 提供统一协议，使模型能够在标准接口下调用文件、数据库、API 或开发工具。对于本系统而言，MCP 可以作为后续扩展方向，让聊天助手从问答系统升级为能够调用工具、读取外部数据并辅助完成任务的 Agent。",
    )

    add_heading(doc, "四、执行结果截图与功能说明")
    add_paragraph(
        doc,
        "本节按照系统搭建和使用顺序记录关键界面。每张截图都对应系统中使用大模型技术的一个环节，能够展示从本地部署、模型接入到知识库问答、工作流编排和工具扩展的完整路径。",
    )

    figure_texts = [
        "Dify 平台启动后，Docker Desktop 中可以看到 web、api、worker、db、redis、nginx 等服务组件处于运行状态。浏览器访问本地地址后进入初始化或登录页面，说明大模型应用平台已经具备运行基础，后续模型配置和应用创建都在该平台中完成。",
        "完成 Ollama 与 DeepSeek 模型安装后，在 Dify 的模型供应商页面中添加本地模型。模型被设置为系统可用模型后，聊天助手、知识库应用和工作流都可以统一调用该模型进行推理。",
        "在 Dify 工作室中创建聊天助手应用，并选择本地 DeepSeek 模型作为推理模型。页面中的模型参数用于控制回答风格、上下文长度和输出范围，用户问题会通过该应用入口传递给大模型处理。",
        "知识库创建时上传本人准备的课程或业务资料，并配置分段、Embedding 和索引方式。系统完成解析与向量化后，文档片段就可以作为后续问答的上下文来源。",
        "将知识库添加到聊天助手后，用户输入与资料相关的问题，系统先检索知识库中的相关片段，再由大模型结合上下文生成回答。该流程体现了 RAG 在系统中的使用方式。",
        "SQL 查询器通过 Dify 工作流实现。画布中开始节点、LLM 节点、参数提取器和 SQL 查询节点依次连接，使自然语言输入能够被转化为可执行的数据查询流程。",
        "参数提取器用于从大模型输出中抽取真正需要执行的 SQL 语句。通过设置参数名称、参数类型和提取指令，系统能够把模型生成的长文本整理成后续数据库组件可以读取的结构化结果。",
        "MCP 或 Agent 工具配置页面展示了大模型连接外部工具和数据源的扩展方向。后续系统可以通过该机制接入 API、数据库或本地工具，使大模型不仅能回答问题，也能调用工具完成任务。",
    ]

    for fig, text in zip(FIGURES, figure_texts):
        add_paragraph(doc, text)
        add_figure(doc, fig)

    add_heading(doc, "五、质量控制与安全考虑")
    add_paragraph(
        doc,
        "大模型应用不仅要能运行，还要关注回答质量和使用安全。知识库问答中，应通过检查引用片段、对比原文内容和多轮测试问题来判断回答是否可靠；SQL 工作流中，应通过提示词约束、参数提取、白名单校验和查询日志降低误执行风险；工具调用场景中，应通过权限控制限制模型能够访问的数据范围。",
    )
    add_paragraph(
        doc,
        "从系统维护角度看，知识库资料需要随着课程内容或业务资料更新而持续维护，提示词也需要根据测试结果不断优化。对于回答不准确、SQL 生成错误或检索片段不相关的情况，应分别从资料质量、分段策略、模型参数和提示词设计四个方面排查。",
    )

    add_heading(doc, "六、总结")
    add_paragraph(
        doc,
        "通过本次“大模型技术”培训与 Dify 平台实践，可以看出大模型在系统中的价值不只体现在文本生成，还体现在自然语言理解、知识检索、流程编排和工具调用等多个方面。Dify 降低了大模型应用开发门槛，Ollama 让本地模型部署更加方便，DeepSeek 模型提供推理与生成能力，知识库和工作流则把模型能力真正接入业务流程。",
    )
    add_paragraph(
        doc,
        "本系统的核心思路是让大模型负责理解与生成，让 Dify 负责组织与调度，让知识库、数据库组件和外部工具负责提供数据与执行能力。这种分工既能体现大模型技术在系统中的应用价值，也便于后续继续扩展更多业务场景。总体来看，本次作业完成了从平台部署、模型接入到应用构建和能力扩展的闭环，较完整地说明了如何在系统中使用大模型技术。",
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_doc())
