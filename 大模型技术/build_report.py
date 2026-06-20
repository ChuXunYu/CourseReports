from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageChops


ROOT = Path(__file__).resolve().parent
RENDER_DIR = ROOT / "_render"
ASSET_DIR = ROOT / "_report_assets"
OUT_DOCX = ROOT / "大模型技术在系统中的应用说明报告.docx"


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size=None, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_style_font(style, size, bold=False):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_keep_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if value:
        if keep_next is None:
            p_pr.append(OxmlElement("w:keepNext"))
    elif keep_next is not None:
        p_pr.remove(keep_next)


def crop_non_white(src: Path, dst: Path, padding=32):
    img = Image.open(src).convert("RGB")
    diff = ImageChops.difference(img, Image.new("RGB", img.size, "white"))
    bbox = diff.getbbox()
    if bbox is None:
        img.save(dst)
        return dst
    left, top, right, bottom = bbox
    left = max(0, left - padding)
    top = max(0, top - padding)
    right = min(img.width, right + padding)
    bottom = min(img.height, bottom + padding)
    img.crop((left, top, right, bottom)).save(dst)
    return dst


def crop_box(src: Path, dst: Path, box):
    img = Image.open(src).convert("RGB")
    img.crop(box).save(dst)
    return dst


def stack_images(srcs, dst: Path, gap=24, bg="white"):
    images = [Image.open(src).convert("RGB") for src in srcs]
    width = max(img.width for img in images)
    height = sum(img.height for img in images) + gap * (len(images) - 1)
    canvas = Image.new("RGB", (width, height), bg)
    y = 0
    for img in images:
        x = (width - img.width) // 2
        canvas.paste(img, (x, y))
        y += img.height + gap
    canvas.save(dst)
    return dst


def add_paragraph(doc, text="", style="Body", indent=True, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Pt(24) if indent else None
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_heading(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = None
    set_paragraph_keep_next(p, True)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 13, bold=True)
    return p


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        set_run_font(run, size=10.5)


def add_caption(doc, caption):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(caption)
    set_run_font(run, size=10.5)
    return p


def add_figure(doc, image_path: Path, caption: str, width_cm=14.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = None
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_metadata_line(doc, label, value):
    p = doc.add_paragraph(style="Meta")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(label)
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(value)
    set_run_font(r2, size=11)


def build_assets():
    ensure_dir(ASSET_DIR)
    result = {}

    p09 = RENDER_DIR / "p09-09.png"
    docker = crop_box(p09, ASSET_DIR / "fig1a_docker.png", (185, 180, 1325, 915))
    install = crop_box(p09, ASSET_DIR / "fig1b_install.png", (185, 1060, 1325, 1595))
    result["fig1"] = stack_images([docker, install], ASSET_DIR / "fig1_local_deploy.png", gap=24)

    result["fig2"] = crop_box(
        RENDER_DIR / "p14-14.png",
        ASSET_DIR / "fig2_model_link.png",
        (160, 150, 1325, 695),
    )

    p16 = RENDER_DIR / "p16-16.png"
    create_entry = crop_box(p16, ASSET_DIR / "fig3a_create_entry.png", (165, 175, 1325, 630))
    chat_type = crop_box(p16, ASSET_DIR / "fig3b_chat_type.png", (165, 730, 1325, 1295))
    result["fig3"] = stack_images([create_entry, chat_type], ASSET_DIR / "fig3_chat_app.png", gap=22)

    p20 = RENDER_DIR / "p20-20.png"
    upload = crop_box(p20, ASSET_DIR / "fig4a_upload.png", (165, 180, 1325, 725))
    embedding = crop_box(p20, ASSET_DIR / "fig4b_embedding.png", (165, 820, 1325, 1430))
    result["fig4"] = stack_images([upload, embedding], ASSET_DIR / "fig4_knowledge_base.png", gap=22)

    result["fig5"] = crop_box(
        RENDER_DIR / "p22-22.png",
        ASSET_DIR / "fig5_rag_app.png",
        (165, 1060, 1325, 1650),
    )

    result["fig6"] = crop_box(
        RENDER_DIR / "p26-26.png",
        ASSET_DIR / "fig6_sql_workflow.png",
        (165, 180, 1325, 730),
    )

    result["fig7"] = crop_box(
        RENDER_DIR / "p29-29.png",
        ASSET_DIR / "fig7_param_extract.png",
        (150, 250, 1325, 820),
    )

    result["fig8"] = crop_box(
        RENDER_DIR / "p35-35.png",
        ASSET_DIR / "fig8_mcp_agent.png",
        (155, 245, 1325, 1435),
    )
    return result


def build_doc():
    assets = build_assets()
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    set_style_font(styles["Normal"], 12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(6)

    body = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(body, 12)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(6)

    meta = styles.add_style("Meta", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(meta, 11)

    code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, 10.5)

    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Caption", 10.5)]:
        set_style_font(styles[name], size, bold=name.startswith("Heading"))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 1.2
    tr = title.add_run("大模型技术在系统中的应用说明报告")
    set_run_font(tr, size=18, bold=True)

    add_metadata_line(doc, "课程内容：", "大模型技术培训")
    add_metadata_line(doc, "培训时间：", "2026 年 6 月 17 日 18:30")
    add_metadata_line(doc, "培训地点：", "信息 A109")
    add_metadata_line(doc, "主讲教师：", "宋航老师")
    add_paragraph(
        doc,
        "本报告围绕“大模型技术”培训任务展开，结合附件《大模型及 Dify 平台介绍》的内容，说明在系统中如何完成大模型能力接入、知识库增强问答、工作流编排以及外部工具扩展。报告重点不只是描述大模型概念，而是把大模型真正放入一个可运行的 Dify 应用流程中，展示从模型部署到业务调用的完整使用方式。",
    )

    add_heading(doc, "一、作业要求与完成情况")
    add_paragraph(
        doc,
        "本次作业要求提交一个文档，说明如何在系统中使用大模型技术。根据培训资料，本系统采用 Dify 作为大模型应用开发平台，以 Ollama 作为本地模型运行环境，并接入 deepseek-r1:1.5b 模型完成基础对话、知识库问答和 SQL 查询工作流等功能。Dify 负责应用编排、知识库管理和流程调度，Ollama 负责在本地提供模型推理服务，两者结合后可以在较低门槛下完成一个可交互、可扩展的大模型应用。",
    )
    add_paragraph(
        doc,
        "从完成情况看，系统围绕三个层次使用大模型技术。第一层是通用对话能力，用户可以在聊天助手中输入自然语言问题，由 DeepSeek 模型生成回答；第二层是知识库增强能力，系统将本地文档切分并向量化，用户提问时先检索相关知识片段，再交给大模型组织答案；第三层是工具调用能力，通过 Dify 工作流把用户输入、LLM 组件、参数提取器和数据库查询组件连接起来，使大模型能够辅助生成 SQL 并调用外部数据源。",
    )
    add_paragraph(
        doc,
        "这种设计符合培训中对大模型应用的要求：大模型不是单独存在的聊天窗口，而是作为系统中的智能推理与自然语言理解模块，与知识库、数据库、工作流和外部工具协同工作，从而让系统具备问答、检索、分析和自动化处理能力。",
    )

    add_heading(doc, "二、系统功能流程说明")
    add_paragraph(
        doc,
        "系统的整体流程可以概括为“本地部署环境、接入模型、创建应用、配置知识库、编排工作流、扩展工具调用”。首先通过 Docker 部署 Dify，并在浏览器中访问本地地址完成管理员初始化；随后安装 Ollama，在本地启动模型服务并下载 deepseek-r1:1.5b 模型；接着在 Dify 中添加 Ollama 模型供应商，把本地模型设置为系统推理模型；最后根据业务场景创建聊天助手、知识库应用或 SQL 查询工作流。",
    )
    add_paragraph(
        doc,
        "在普通聊天场景中，用户输入自然语言问题后，Dify 将问题发送给已配置的大模型，模型根据提示词、上下文和参数设置生成回复。该流程适合基础答疑、内容生成、学习辅助和通用咨询等任务。",
    )
    add_paragraph(
        doc,
        "在知识库问答场景中，系统先把上传的 PDF、Markdown、Word 等资料解析为文本片段，再通过 Embedding 模型转成向量并保存到知识库。用户提问时，系统根据问题向量检索最相关的文档片段，将检索结果与用户问题共同传给大模型生成答案。相比单纯调用模型，这种方式可以补充模型训练后没有的新知识，也能让回答更贴近课程资料或业务文档。",
    )
    add_paragraph(
        doc,
        "在 SQL 查询场景中，Dify 工作流把“大模型理解需求”和“数据库组件执行查询”分开处理。用户先用自然语言描述查询目标，LLM 组件根据表名、字段说明和输出格式要求生成 SQL，参数提取器再从模型输出中提取结构化 SQL 数组，最后由数据库查询组件执行并返回结果。这个流程体现了大模型在系统中的作用：它负责理解、规划和生成，而具体数据访问仍由受控组件执行。",
    )
    add_paragraph(
        doc,
        "在进一步扩展场景中，MCP 可以作为模型连接外部工具和数据源的标准协议。通过 MCP，大模型能够在权限控制下访问数据库、文件系统、API 或开发工具，使系统从“被动回答问题”扩展为“主动调用工具完成任务”的 Agent 形态。",
    )

    add_heading(doc, "三、系统设计与实现要点")
    add_heading(doc, "3.1 Dify 与本地模型的连接", level=2)
    add_paragraph(
        doc,
        "Dify 通过 Docker 运行，而 Ollama 运行在本地电脑上，因此系统需要让容器内的 Dify 能够访问本机的 Ollama 服务。培训资料中采用在 Dify 的 docker 配置文件末尾追加环境变量的方式开启自定义模型，并把 Ollama 地址设置为 host.docker.internal:11434。这样 Dify 在调用模型时就可以通过本地 API 访问 Ollama 中运行的 DeepSeek 模型。",
    )
    add_code(
        doc,
        [
            "# 启用自定义模型",
            "CUSTOM_MODEL_ENABLED=true",
            "# 指定 Ollama 的 API 地址",
            "OLLAMA_API_BASE_URL=host.docker.internal:11434",
        ],
    )
    add_paragraph(
        doc,
        "完成环境配置后，在 Dify 的“设置—模型供应商—Ollama”中添加模型，并在系统模型设置里选择 deepseek-r1:1.5b。这个步骤完成后，Dify 应用层就能统一调用本地模型，后续聊天助手、知识库和工作流都可以复用同一个模型能力。",
    )

    add_heading(doc, "3.2 对话应用中的提示词与参数配置", level=2)
    add_paragraph(
        doc,
        "创建聊天助手时，系统需要选择合适的模型并设置推理参数，例如温度、Top P、上下文长度和最大输出长度。参数配置会影响回答的稳定性和创造性：学习答疑、知识库问答和 SQL 生成更适合偏稳定的参数，避免模型输出过于发散；开放式讨论或内容创作可以适当提高创造性，使回答更丰富。",
    )
    add_paragraph(
        doc,
        "提示词用于约束模型身份、任务边界和输出格式。在本系统中，提示词不只是“告诉模型回答问题”，还会说明数据库类型、表结构、字段含义、返回 JSON 的格式要求以及不能使用的 SQL 语法。通过这种方式，大模型输出会更加可控，方便后续组件继续处理。",
    )

    add_heading(doc, "3.3 知识库增强生成的实现", level=2)
    add_paragraph(
        doc,
        "知识库部分使用了 RAG 的思路。系统把课程资料或业务文档上传到 Dify 知识库后，会进行文档解析、分段、Embedding 向量化和索引保存。用户提问时，系统不是只依赖模型自身记忆，而是先检索知识库中最相关的片段，再让模型结合这些片段组织答案。这样可以降低“模型一本正经地编答案”的风险，也能让回答保持在指定资料范围内。",
    )
    add_paragraph(
        doc,
        "从实际应用角度看，知识库适合承载课程讲义、项目文档、制度说明、产品手册和数据库字段解释等材料。模型负责理解问题和归纳表达，知识库负责提供可信上下文，二者结合后系统能回答更具体的专业问题。",
    )

    add_heading(doc, "3.4 SQL 工作流的编排", level=2)
    add_paragraph(
        doc,
        "SQL 查询工作流体现了 Dify 的可视化编排能力。开始节点接收用户输入，LLM 节点根据提示词生成查询方案，参数提取器从模型输出中提取 SQL 数组，数据库查询组件再执行查询并返回结果。为了让流程稳定运行，提示词中明确要求使用 MySQL 5.7 语法，避免使用不兼容函数，并要求输出能够被参数提取器识别的结构化内容。",
    )
    add_paragraph(
        doc,
        "在真实系统中，还需要增加安全控制，例如限制只允许 SELECT 查询、禁止 DROP、DELETE、UPDATE 等危险语句，对查询字段和表名做白名单校验，并记录查询日志。这样既能发挥大模型理解自然语言的优势，又能避免模型直接操作数据库带来的风险。",
    )

    add_heading(doc, "3.5 MCP 与 Agent 扩展", level=2)
    add_paragraph(
        doc,
        "MCP（Model Context Protocol，模型上下文协议）解决的是大模型连接外部数据源和工具的问题。传统做法通常需要为每个模型和每个工具分别写接口，而 MCP 通过统一协议把模型、工具和数据源连接起来。对于本系统而言，MCP 可以作为后续升级方向，使模型在权限控制下调用天气 API、企业数据库、本地文件或开发工具，从而进一步构建具备行动能力的 Agent。",
    )

    add_heading(doc, "四、执行结果截图与功能说明")
    add_paragraph(
        doc,
        "本节结合培训附件中的操作界面，对系统中使用大模型技术的关键环节进行说明。截图按照实际搭建顺序排列，展示从本地部署到模型连接、从知识库构建到工作流执行的完整过程。",
    )

    add_paragraph(
        doc,
        "在 Docker Desktop 中可以看到 Dify 相关容器已经运行，包括 web、api、worker、db、redis、nginx 等组件。浏览器访问 http://localhost/install 后进入管理员账号初始化页面，说明 Dify 平台已经在本地启动，后续模型接入和应用创建都在该平台中完成。",
    )
    add_figure(doc, assets["fig1"], "图 1 本地部署 Dify 并初始化管理员账号", width_cm=12.8)

    add_paragraph(
        doc,
        "完成 Ollama 和 DeepSeek 模型安装后，在 Dify 的模型供应商页面可以看到系统推理模型选择为 deepseek-r1:1.5b。此时 Dify 与本地模型服务已经连通，应用中的聊天助手、知识库和工作流都可以调用该模型进行推理。",
    )
    add_figure(doc, assets["fig2"], "图 2 在 Dify 中关联本地 DeepSeek 模型", width_cm=10.8)

    add_paragraph(
        doc,
        "进入工作室后点击“创建空白应用”，选择“聊天助手”，输入应用名称和描述后创建应用。创建完成后在右上角选择模型并设置温度、Top P、上下文长度等参数。这个页面是系统使用大模型的第一个入口，用户的问题会通过聊天助手传递给模型并得到自然语言回复。",
    )
    add_figure(doc, assets["fig3"], "图 3 创建聊天助手并配置模型参数", width_cm=13.5)

    add_paragraph(
        doc,
        "知识库创建时，系统支持导入已有文本或上传 PDF、DOCX、CSV 等文件。上传资料后选择 Embedding 模型和检索方式，Dify 会自动完成文档解析、分段和向量化处理。处理完成后，知识库中的片段可以作为后续问答的上下文。",
    )
    add_figure(doc, assets["fig4"], "图 4 上传资料并完成知识库向量化配置", width_cm=13.2)

    add_paragraph(
        doc,
        "将知识库添加到聊天助手后，用户在对话框中输入问题，系统会先从知识库中检索相关片段，再由大模型根据检索内容生成回答。这样构建出的知识库应用可以回答课程资料或业务资料中的专业问题，弥补基础模型不了解本地文档内容的不足。",
    )
    add_figure(doc, assets["fig5"], "图 5 将知识库添加到聊天应用并进行问答测试", width_cm=13.2)

    add_paragraph(
        doc,
        "SQL 查询器通过工作流实现。页面中可以看到开始节点、LLM 节点、参数提取器和 SQL 查询节点之间的连接关系。开始节点负责接收用户输入，LLM 节点负责理解查询意图并生成 SQL，后续节点负责提取和执行，从而把自然语言查询转化为数据库操作。",
    )
    add_figure(doc, assets["fig6"], "图 6 SQL 查询工作流的节点编排", width_cm=13.4)

    add_paragraph(
        doc,
        "由于 DeepSeek 模型会输出思考和解释内容，系统需要通过参数提取器从模型返回结果中取出真正可执行的 SQL 语句。参数提取器要求输出为数组格式，每个元素包含一条 SQL，这样数据库查询组件才能稳定读取并执行。",
    )
    add_figure(doc, assets["fig7"], "图 7 使用参数提取器提取 SQL 语句", width_cm=13.2)

    add_paragraph(
        doc,
        "MCP 部分说明了大模型如何连接外部数据源和工具。大模型本身无法直接访问实时数据、本地文件或数据库，MCP 通过标准协议让模型能够在统一接口下调用工具，解决数据孤岛、重复开发和生态碎片化问题。后续如果将 MCP 服务接入系统，聊天助手就可以进一步具备查询外部 API、读取文件和调用开发工具的能力。",
    )
    add_figure(doc, assets["fig8"], "图 8 MCP 用于连接外部工具与数据源", width_cm=12.0)

    add_heading(doc, "五、总结")
    add_paragraph(
        doc,
        "通过本次“大模型技术”培训与 Dify 平台实践，可以看出大模型在系统中的价值不只体现在生成文本，还体现在自然语言理解、知识检索、流程编排和工具调用等多个方面。Dify 降低了大模型应用开发门槛，Ollama 让本地模型部署更加方便，DeepSeek 模型提供推理与生成能力，知识库和工作流则把模型能力真正接入业务流程。",
    )
    add_paragraph(
        doc,
        "本系统的核心思路是让大模型承担“理解和生成”的工作，让 Dify 负责“组织和调度”，让知识库、数据库组件和 MCP 工具负责“提供可靠数据与执行能力”。这种分工可以提高系统可用性，也便于后续扩展更多业务功能。后续优化方向包括完善提示词模板、加强 SQL 安全校验、记录模型调用日志、评估回答准确率，并根据实际业务资料持续更新知识库。",
    )
    add_paragraph(
        doc,
        "总体来说，本次作业完成了从理论理解到系统落地的闭环：先认识大模型的定义、特点和应用场景，再通过 Dify 平台把模型接入实际应用，最后利用知识库、工作流和 MCP 思路扩展系统能力。这样的实现方式能够较完整地体现“大模型技术在系统中的使用”。",
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    out = build_doc()
    print(out)
