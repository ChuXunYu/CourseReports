from pathlib import Path
import textwrap

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT = Path(__file__).resolve().parent
SCREENSHOT_DIR = ROOT / "report_screenshots"
ASSET_DIR = ROOT / "_report_assets"
OUT_DOCX = ROOT / "鸿蒙开发技术在教学设备资产管理系统中的应用说明报告_待替换截图.docx"


FIGURES = [
    {
        "key": "fig01",
        "file": "fig01_deveco_project.png",
        "caption": "图 1 DevEco Studio 中的鸿蒙工程结构",
        "title": "DevEco Studio 中的鸿蒙工程结构",
        "requirements": ["打开 HarmonyAssetModule 工程", "显示 entry、ets、pages、common、module.json5 等目录或文件", "能看出这是 HarmonyOS/ArkTS 工程"],
        "width_cm": 13.3,
    },
    {
        "key": "fig02",
        "file": "fig02_arkts_index_code.png",
        "caption": "图 2 使用 ArkTS 与 ArkUI 编写资产列表页面",
        "title": "使用 ArkTS 与 ArkUI 编写资产列表页面",
        "requirements": ["打开 entry/src/main/ets/pages/Index.ets", "能看到 Navigation、List、TextInput、ForEach 等 ArkUI 组件", "能看到搜索、状态筛选或资产卡片代码"],
        "width_cm": 13.3,
    },
    {
        "key": "fig03",
        "file": "fig03_home_preview.png",
        "caption": "图 3 资产台账鸿蒙端首页预览",
        "title": "资产台账鸿蒙端首页预览",
        "requirements": ["打开 DevEco Studio Previewer 或模拟器", "显示教学设备资产首页", "能看到资产台账、搜索框、状态筛选和资产卡片"],
        "width_cm": 12.2,
    },
    {
        "key": "fig04",
        "file": "fig04_search_filter.png",
        "caption": "图 4 鸿蒙端资产搜索与状态筛选",
        "title": "鸿蒙端资产搜索与状态筛选",
        "requirements": ["在搜索框输入关键字或点击状态筛选", "列表结果发生变化", "截图中同时出现筛选条件和筛选后的资产列表"],
        "width_cm": 12.2,
    },
    {
        "key": "fig05",
        "file": "fig05_asset_detail.png",
        "caption": "图 5 资产详情页面",
        "title": "资产详情页面",
        "requirements": ["点击资产卡片进入详情页", "显示资产编号、分类、部门、供应商、保管人、资产原值和备注", "能看到返回或状态更新按钮"],
        "width_cm": 12.2,
    },
    {
        "key": "fig06",
        "file": "fig06_status_update.png",
        "caption": "图 6 移动端资产状态更新交互",
        "title": "移动端资产状态更新交互",
        "requirements": ["在详情页点击标记为使用中", "页面状态文字更新为使用中", "截图中应能看到操作后的状态变化"],
        "width_cm": 12.2,
    },
    {
        "key": "fig07",
        "file": "fig07_multi_device_preview.png",
        "caption": "图 7 手机和平板预览体现一次开发多端部署",
        "title": "手机和平板预览体现一次开发多端部署",
        "requirements": ["在 DevEco Studio 切换不同设备预览", "至少展示手机预览；如果可以，拼接平板预览", "能体现同一 ArkTS 页面适配不同屏幕"],
        "width_cm": 13.0,
    },
    {
        "key": "fig08",
        "file": "fig08_widget_design.png",
        "caption": "图 8 万能卡片或元服务扩展设计",
        "title": "万能卡片或元服务扩展设计",
        "requirements": ["可以展示卡片设计草图、DevEco 卡片模板或报告中的卡片原型", "卡片内容建议包含资产总数、维修中数量、快速进入按钮", "如暂未实现卡片，可截 DevEco 中卡片页面/设计说明"],
        "width_cm": 12.4,
    },
]


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def font_path(name: str) -> str | None:
    path = Path("C:/Windows/Fonts") / name
    return str(path) if path.exists() else None


def load_font(size: int, bold: bool = False):
    names = ["msyhbd.ttc", "simhei.ttf", "msyh.ttc", "simsun.ttc"] if bold else ["msyh.ttc", "simsun.ttc", "simhei.ttf"]
    for name in names:
        path = font_path(name)
        if path:
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_text(text: str, width: int) -> list[str]:
    return textwrap.wrap(text, width=width, break_long_words=False, replace_whitespace=False)


def create_placeholder(fig: dict) -> Path:
    ensure_dir(ASSET_DIR)
    out = ASSET_DIR / f"{fig['key']}_placeholder.png"
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((52, 52, 1548, 848), radius=10, fill=(246, 248, 252), outline=(142, 148, 158), width=3)
    draw.rectangle((52, 52, 1548, 166), fill=(232, 238, 248))
    draw.line((52, 166, 1548, 166), fill=(142, 148, 158), width=2)
    title_font = load_font(42, True)
    sub_font = load_font(27, True)
    body_font = load_font(28)
    small_font = load_font(23)
    draw.text((92, 84), "待替换为鸿蒙实操截图", font=title_font, fill=(37, 76, 126))
    draw.text((92, 206), fig["caption"], font=sub_font, fill=(18, 24, 32))
    draw.text((92, 256), f"建议文件名：report_screenshots/{fig['file']}", font=small_font, fill=(92, 99, 105))
    draw.text((92, 330), "截图必须包含：", font=sub_font, fill=(18, 24, 32))
    y = 386
    for i, req in enumerate(fig["requirements"], start=1):
      for line in wrap_text(f"{i}. {req}", 42):
        draw.text((120, y), line, font=body_font, fill=(30, 36, 42))
        y += 43
      y += 8
    draw.text((92, 760), "说明：该图仅用于占位，最终提交前请用 DevEco Studio 或模拟器真实截图替换。", font=small_font, fill=(126, 74, 0))
    img.save(out)
    return out


def screenshot_path(fig: dict) -> Path:
    if SCREENSHOT_DIR.exists():
        primary = SCREENSHOT_DIR / fig["file"]
        if primary.exists():
            return primary
        stem = Path(fig["file"]).stem
        for ext in (".png", ".jpg", ".jpeg"):
            candidate = SCREENSHOT_DIR / f"{stem}{ext}"
            if candidate.exists():
                return candidate
    return create_placeholder(fig)


def normalize_image(path: Path, key: str) -> Path:
    ensure_dir(ASSET_DIR)
    out = ASSET_DIR / f"{key}_normalized.png"
    with Image.open(path) as img:
        img = ImageOps.exif_transpose(img).convert("RGB")
        img.thumbnail((1800, 1400), Image.LANCZOS)
        canvas = Image.new("RGB", img.size, "white")
        canvas.paste(img, (0, 0))
        canvas.save(out, optimize=True)
    return out


def set_run_font(run, size=None, bold=False):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_style_font(style, size, bold=False):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def keep_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def add_paragraph(doc, text="", style="Body", indent=True, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Pt(24) if indent else None
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 12)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style="Heading 1" if level == 1 else "Heading 2")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    keep_next(p)
    run = p.add_run(text)
    set_run_font(run, 15 if level == 1 else 13, True)


def add_meta(doc, label, value):
    p = doc.add_paragraph(style="Meta")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label)
    set_run_font(r1, 11, True)
    r2 = p.add_run(value)
    set_run_font(r2, 11)


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        set_run_font(run, 10.5)


def add_caption(doc, caption):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(caption)
    set_run_font(run, 10.5)


def add_figure(doc, fig):
    img = normalize_image(screenshot_path(fig), fig["key"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(img), width=Cm(fig["width_cm"]))
    add_caption(doc, fig["caption"])


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    styles = doc.styles
    set_style_font(styles["Normal"], 12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(6)
    body = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(body, 12)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(6)
    meta = styles.add_style("Meta", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(meta, 11)
    code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, 10.5)
    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Caption", 10.5)]:
        set_style_font(styles[name], size, name.startswith("Heading"))
    return doc


def build_doc():
    doc = setup_doc()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("鸿蒙开发技术在教学设备资产管理系统中的应用说明报告")
    set_run_font(run, 18, True)

    add_meta(doc, "课程主题：", "鸿蒙开发技术")
    add_meta(doc, "项目名称：", "教学设备资产管理系统鸿蒙端模块")
    add_meta(doc, "开发形态：", "HarmonyOS 手机端应用模块，可扩展为元服务与万能卡片")
    add_meta(doc, "实现方式：", "DevEco Studio + HarmonyOS SDK + ArkTS + ArkUI")
    add_meta(doc, "核心功能：", "资产列表、搜索筛选、详情查看、资产状态更新")
    add_paragraph(doc, "本报告围绕鸿蒙开发技术作业要求展开，选择已开发系统中的核心功能模块进行鸿蒙化实现。原系统的核心业务是教学设备资产台账管理，本次将资产列表、搜索筛选、详情查看和状态更新等功能迁移为 HarmonyOS 移动端模块，并说明其工程结构、页面实现、交互流程以及后续元服务和万能卡片扩展思路。")
    add_paragraph(doc, "报告中的鸿蒙端实现不是对原 Web 页面做简单截图说明，而是将资产管理中的高频业务重新组织为适合移动端和多设备场景的入口。资产管理员在实训室现场可以直接通过手机查看资产状态、定位设备信息并完成轻量级维护；在后续扩展中，还可以通过元服务或万能卡片把常用信息前置到桌面，实现更短的操作路径。")

    add_heading(doc, "一、作业要求与完成情况")
    add_paragraph(doc, "本次任务要求选取所开发系统的核心功能模块，实现鸿蒙化开发，可以是 PC 端应用、手机端应用、平板端应用或其他嵌入式设备应用。根据该要求，本项目选取“教学设备资产管理系统”的资产台账模块，使用 ArkTS 和 ArkUI 设计 HarmonyOS 手机端页面，使用户能够在移动端查看资产列表、搜索资产、按状态筛选、进入详情页并更新资产状态。")
    add_paragraph(doc, "之所以选择资产台账模块，是因为它连接资产数据、现场盘点、状态维护和后续统计，是系统中最能体现业务价值的核心流程。若只展示登录或菜单页面，难以说明鸿蒙端对实际业务的改造效果；选择资产台账后，可以同时呈现数据列表、触摸筛选、详情查看、状态变更和多端适配等关键能力。")
    add_paragraph(doc, "从完成情况看，鸿蒙端模块覆盖了原 Web 系统中最核心的资产查询与维护流程，并围绕移动端使用习惯进行了界面重组。首页将搜索框、状态筛选和资产卡片放在同一页面，减少用户操作层级；详情页集中展示资产编号、分类、部门、供应商、保管人、资产原值和备注，并提供状态更新入口。该模块既能作为独立手机端应用运行，也可以继续拆分为元服务或万能卡片。")
    add_paragraph(doc, "本次实现体现了鸿蒙开发的几个核心特点：一是使用 DevEco Studio 和 HarmonyOS SDK 进行工程管理、页面预览和调试；二是使用 ArkTS 声明式语法构建界面；三是通过 ArkUI 组件实现移动端交互；四是围绕一次开发、多端部署的思路预留手机和平板适配空间；五是考虑万能卡片将资产统计和快捷入口前置到桌面。")

    add_heading(doc, "二、系统功能流程说明")
    add_paragraph(doc, "鸿蒙端模块的业务流程从资产列表进入。用户打开应用后进入“教学设备资产”首页，页面显示资产台账标题、当前筛选结果数量、搜索输入框、状态筛选按钮和资产卡片列表。资产卡片展示资产名称、资产编号、状态、分类、使用部门、保管人和资产原值，用户不需要进入详情页就可以快速判断资产基本情况。")
    add_paragraph(doc, "当用户需要查找某个资产时，可以在搜索框输入资产编号、名称、部门或保管人，页面会根据输入内容实时过滤列表。用户也可以点击“全部”“在库”“使用中”“维修中”等状态按钮，快速查看不同状态下的设备。搜索和筛选都在移动端页面中完成，适合资产管理员在实训室巡检、盘点或维修登记时使用。")
    add_paragraph(doc, "用户点击某张资产卡片后进入资产详情页面。详情页以信息块方式展示当前状态、设备分类、使用部门、供应商、保管人、资产原值和备注等字段。页面底部提供“标记为使用中”和“返回资产列表”两个操作按钮，表示移动端不仅用于浏览数据，也可以承接轻量级业务维护。后续对接后端接口后，该状态更新操作可以同步写回原有资产管理系统。")

    add_heading(doc, "三、系统设计与实现要点")
    add_heading(doc, "3.1 鸿蒙技术在本系统中的落点", level=2)
    add_paragraph(doc, "在本系统中，HarmonyOS 应用形态承担移动端资产管理入口。用户在手机或平板上打开应用后，直接进入资产台账页面，不再依赖浏览器访问 Web 后台。该入口适合实训室资产巡检、设备借用登记、维修状态查看等现场业务。")
    add_paragraph(doc, "元服务形态适合作为资产查询或维修登记的轻量入口。后续可以将“扫码查看资产详情”“提交维修状态”“查看待盘点资产”等功能拆成独立元服务，使用户在扫一扫、碰一碰或系统推荐入口中直达具体服务。万能卡片则适合展示资产总数、维修中数量、待盘点数量和快捷进入按钮，把资产管理中的高频信息放在桌面或负一屏。")
    add_paragraph(doc, "一次开发、多端部署体现在页面组件和布局策略上。当前模块使用 ArkUI 的 Column、Row、List 等基础组件组织页面，核心业务逻辑保存在 ArkTS 状态和数据模型中。手机端以单列卡片列表为主，平板端可以在同一业务组件基础上扩展为列表与详情分栏，减少重复开发。")

    add_heading(doc, "3.2 HarmonyOS 工程结构", level=2)
    add_paragraph(doc, "鸿蒙模块采用 Stage 模型组织工程，主要目录包括 AppScope、entry、entry/src/main/ets、resources 和 module.json5。AppScope 中保存应用级配置，entry 模块作为主入口，module.json5 描述模块类型、支持设备和 EntryAbility，ets/pages 存放页面文件，ets/common 存放资产数据模型。这样的结构符合 HarmonyOS 应用工程的基本组织方式。")
    add_code(doc, [
        "HarmonyAssetModule/",
        "  AppScope/app.json5",
        "  entry/src/main/module.json5",
        "  entry/src/main/ets/entryability/EntryAbility.ets",
        "  entry/src/main/ets/pages/Index.ets",
        "  entry/src/main/ets/pages/AssetDetail.ets",
        "  entry/src/main/ets/common/AssetModel.ets",
    ])
    add_paragraph(doc, "EntryAbility 负责应用窗口创建和首页加载，Index.ets 负责资产列表和筛选，AssetDetail.ets 负责详情展示和状态更新，AssetModel.ets 定义资产数据结构和状态显示映射。通过这种分层方式，页面展示、业务数据和应用入口相互分离，便于后续扩展接口请求、缓存或卡片能力。")

    add_heading(doc, "3.3 ArkTS 与 ArkUI 页面实现", level=2)
    add_paragraph(doc, "页面使用 ArkTS 声明式开发方式实现。Index 页面通过 @State 保存关键字和状态筛选条件，通过计算属性 filteredAssets 根据用户输入动态生成列表数据。界面使用 Navigation、Column、Row、TextInput、List、ListItem、ForEach、Button 等 ArkUI 组件组合而成，能够较自然地适配手机屏幕上的纵向浏览和触摸操作。")
    add_code(doc, [
        "@State keyword: string = '';",
        "@State selectedStatus: string = 'ALL';",
        "List({ space: 12 }) {",
        "  ForEach(this.filteredAssets, (asset: AssetItem) => {",
        "    ListItem() { this.AssetCard(asset) }",
        "  }, (asset: AssetItem) => asset.assetCode)",
        "}",
    ])
    add_paragraph(doc, "资产卡片使用圆角、阴影和状态标签突出关键业务信息。详情页使用多个信息块展示资产属性，并使用按钮触发状态更新。这样的界面结构比直接照搬 Web 表格更适合移动端，因为它减少了横向滚动，也让用户在小屏幕上更容易阅读和操作。")

    add_heading(doc, "3.4 一次开发多端部署与元服务扩展", level=2)
    add_paragraph(doc, "HarmonyOS 强调一次开发、多端部署。资产台账页面采用相对弹性的布局，核心内容由 Column、Row 和 List 组成，能够在手机和平板预览中保持相同业务结构。后续如果需要适配平板，可以将列表与详情改为左右分栏；如果需要适配桌面或车机设备，则可以保留资产卡片和状态筛选的业务组件，再调整外层布局。")
    add_paragraph(doc, "该模块也适合扩展为元服务。资产管理员在设备巡检、维修登记或盘点时，往往只需要快速查看某条资产状态，不一定需要完整安装复杂应用。将资产查询或维修登记拆分为元服务后，用户可以通过扫码、碰一碰或入口卡片直接打开对应功能。万能卡片则可以把“资产总数、维修中数量、待盘点数量、快速进入”前置到桌面或负一屏，实现服务直达。")

    add_heading(doc, "3.5 与原系统的数据衔接", level=2)
    add_paragraph(doc, "当前示例模块使用本地 seed 数据模拟资产台账，便于在 DevEco Studio 中快速预览。真实接入时，可以通过 @ohos.net.http 调用后端接口，例如 GET /api/assets 获取资产列表，GET /api/assets/{id} 获取详情，PUT /api/assets/{id}/status 更新状态。鸿蒙端只负责移动端交互和展示，原系统继续负责数据库、权限和业务校验。")
    add_paragraph(doc, "这种设计能保持系统边界清晰：数据库和核心业务规则仍在后端系统中维护，HarmonyOS 模块负责提供更便捷的移动端入口；万能卡片或元服务负责把高频信息前置，减少用户打开完整应用的次数。")

    add_heading(doc, "四、执行结果截图与功能说明")
    add_paragraph(doc, "本节截图按照工程结构、页面代码、首页预览、搜索筛选、详情页面、状态更新、多端预览和卡片扩展的顺序排列。最终提交前需要将占位图替换为队员在 DevEco Studio、Previewer 或模拟器中的真实截图，截图文件名与 picture_capture.md 保持一致。")
    desc = [
        "在 DevEco Studio 中打开工程后，可以看到 entry 模块、ets 页面目录、common 数据模型目录以及 module.json5 等文件。该结构体现了 HarmonyOS Stage 模型工程的基本组成。",
        "Index.ets 页面使用 ArkTS 编写资产列表界面，代码中包含 Navigation、TextInput、List、ForEach 等 ArkUI 组件。页面通过 @State 保存搜索关键字和状态筛选条件，能够根据用户操作实时更新列表。",
        "首页预览展示资产台账移动端界面。页面上方显示资产台账标题和记录数量，下方提供搜索框、状态筛选按钮和资产卡片列表，适合在手机端快速查看设备资产。",
        "用户在搜索框输入关键字或点击状态筛选按钮后，列表会显示符合条件的资产。该交互对应原 Web 系统中的查询功能，但在鸿蒙端采用更适合触屏的按钮和卡片形式。",
        "点击资产卡片后进入详情页。详情页集中展示资产状态、分类、部门、供应商、保管人、原值和备注等信息，让资产管理员能够在移动端快速确认设备情况。",
        "在详情页点击“标记为使用中”后，页面状态文字发生变化，表示移动端可以承接轻量级状态维护。后续对接后端接口后，该操作可以同步写回资产管理系统。",
        "通过 DevEco Studio 切换手机和平板预览，可以观察同一 ArkTS 页面在不同屏幕上的展示效果。这体现了鸿蒙应用一次开发、多端部署的开发思路。",
        "万能卡片或元服务扩展可以将资产总数、维修中数量和快速入口前置到桌面或负一屏。对于资产管理员来说，高频信息不需要进入完整应用即可查看。",
    ]
    for fig, text in zip(FIGURES, desc):
        add_paragraph(doc, text)
        add_figure(doc, fig)

    add_heading(doc, "五、测试与运行说明")
    add_paragraph(doc, "本模块需要在 DevEco Studio 中打开或复制到 HarmonyOS 工程中运行。队员可以新建 Empty Ability 工程，将 Index.ets、AssetDetail.ets 和 AssetModel.ets 放入对应目录，并确认 module.json5 中 pages 配置包含 pages/Index 与 pages/AssetDetail。完成后可以使用 Previewer 查看页面，也可以使用模拟器运行应用。")
    add_paragraph(doc, "由于当前环境未安装 DevEco Studio 和 HarmonyOS SDK，报告中的截图位置先使用占位图。最终提交前，应由安装了 DevEco Studio 的队员根据 picture_capture.md 完成真实截图。截图应覆盖工程结构、代码实现、页面预览和交互结果，使报告能够同时说明鸿蒙技术使用过程和业务功能完成情况。")

    add_heading(doc, "六、总结")
    add_paragraph(doc, "通过本次鸿蒙化开发实践，可以看出 HarmonyOS 不只是把原有 Web 页面搬到手机上，而是要求根据移动端和多设备场景重新组织业务入口。资产管理模块在鸿蒙端使用 ArkTS 和 ArkUI 构建，围绕资产查询、筛选、详情查看和状态维护形成闭环，同时保留了与后端系统对接的接口扩展空间。")
    add_paragraph(doc, "本项目选取教学设备资产管理系统的核心模块进行鸿蒙化，实现了手机端应用页面，并进一步分析了元服务和万能卡片扩展方向。整体上，报告覆盖了 HarmonyOS 工程结构、ArkTS 页面开发、ArkUI 组件使用、移动端交互、多端适配和服务直达设计，能够较完整地说明如何在系统中使用鸿蒙开发技术。")

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_doc())
